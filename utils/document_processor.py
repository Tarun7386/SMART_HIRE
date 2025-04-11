import os
import re

def extract_text_from_file(file_path):
    """Extract text from a file based on its extension"""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext in ['.txt', '.md', '.rtf']:  # This line already includes .txt
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file"""
    try:
        import PyPDF2
        
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        return text
    except ImportError:
        print("PyPDF2 is required to extract text from PDF files. Install it with: pip install PyPDF2")
        # Return file path as a fallback
        return f"PDF file: {os.path.basename(pdf_path)}"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file"""
    try:
        import docx
        
        doc = docx.Document(docx_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    except ImportError:
        print("python-docx is required to extract text from DOCX files. Install it with: pip install python-docx")
        # Return file path as a fallback
        return f"DOCX file: {os.path.basename(docx_path)}"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_from_txt(txt_path):
    """Extract text from a plain text file"""
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        try:
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text with latin-1 encoding: {e}")
            return ""
    except Exception as e:
        print(f"Error extracting text from text file: {e}")
        return ""